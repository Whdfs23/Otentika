"""
document_parser.py
Ekstraksi & pembersihan teks dari PDF maupun DOCX (Word) -- dua format paling
umum dipakai mahasiswa untuk submit tugas. Output akhirnya seragam: teks bersih
siap masuk ke similarity_engine.py, apapun format aslinya.

Install:
    pip install pdfplumber python-docx

Pakai:
    python document_parser.py "path/ke/file.pdf"
    python document_parser.py "path/ke/file.docx"
"""

import re
import sys
from pathlib import Path
from collections import Counter

import pdfplumber
from docx import Document


# ---------- PDF ----------

def _extract_page_zones(path: str, margin_ratio: float = 0.08) -> list[dict]:
    """
    Pisahkan tiap halaman jadi 3 zona berdasarkan POSISI FISIK, bukan tebakan:
    header (~8% atas), footer (~8% bawah), body (sisanya / tengah).
    Cuma header & footer yang nanti dicek pengulangannya -- body tidak pernah
    dianggap noise, sekalipun ada kalimat yang kebetulan berulang di sana
    (misal label struktur konten seperti "Kondisi sebelum perubahan:").
    """
    zones = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            h, w = page.height, page.width
            header = page.crop((0, 0, w, h * margin_ratio)).extract_text() or ""
            footer = page.crop((0, h * (1 - margin_ratio), w, h)).extract_text() or ""
            body = page.crop((0, h * margin_ratio, w, h * (1 - margin_ratio))).extract_text() or ""
            zones.append({"header": header, "footer": footer, "body": body})
    return zones


def _detect_repeating_lines(texts: list[str], min_occurrence_ratio: float = 0.6) -> set[str]:
    line_counter = Counter()
    for text in texts:
        lines = {line.strip() for line in text.split("\n") if line.strip()}
        line_counter.update(lines)
    threshold = max(2, int(len(texts) * min_occurrence_ratio))
    return {line for line, count in line_counter.items() if count >= threshold}


def extract_clean_text_pdf(path: str) -> str:
    zones = _extract_page_zones(path)

    repeating_boilerplate = _detect_repeating_lines(
        [z["header"] for z in zones] + [z["footer"] for z in zones]
    )

    cleaned_lines = []
    for z in zones:
        # Header & footer: buang yang memang berulang (boilerplate asli)
        for zone_name in ("header", "footer"):
            for line in z[zone_name].split("\n"):
                stripped = line.strip()
                if not stripped or stripped in repeating_boilerplate:
                    continue
                if re.fullmatch(r"\d+", stripped):
                    continue
                cleaned_lines.append(stripped)

        # Body: SELALU disimpan utuh, apa pun isinya, walau ada label yang berulang
        for line in z["body"].split("\n"):
            stripped = line.strip()
            if stripped and not re.fullmatch(r"\d+", stripped):
                cleaned_lines.append(stripped)

    return " ".join(cleaned_lines)


# ---------- DOCX ----------

def extract_clean_text_docx(path: str) -> str:
    """
    Word menyimpan header/footer di bagian terpisah dari body (beda dengan PDF
    yang nyatu jadi teks polos), jadi doc.paragraphs otomatis sudah bersih dari
    header/footer berulang. Kita tetap filter baris yang isinya cuma angka
    (nomor halaman manual) untuk jaga-jaga, dan ikut tarik teks dari tabel
    karena laporan tugas sering pakai tabel.
    """
    doc = Document(path)
    lines = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or re.fullmatch(r"\d+", text):
            continue
        lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return " ".join(lines)


# ---------- Dispatcher ----------

def extract_clean_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_clean_text_pdf(path)
    elif ext == ".docx":
        return extract_clean_text_docx(path)
    else:
        raise ValueError(f"Format belum didukung: {ext}. Pakai .pdf atau .docx.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print('Pakai: python document_parser.py "path/ke/file.pdf-atau-.docx"')
        sys.exit(1)

    path = sys.argv[1]
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        zones = _extract_page_zones(path)
        boilerplate = _detect_repeating_lines(
            [z["header"] for z in zones] + [z["footer"] for z in zones]
        )
        print(f"Jumlah halaman: {len(zones)}")
        print(f"Baris di zona header/footer yang dideteksi berulang ({len(boilerplate)}):")
        for line in list(boilerplate)[:10]:
            print(f"  - {line!r}")
        print()

    cleaned = extract_clean_text(path)

    print(f"Format terdeteksi: {ext}")
    print(f"Total karakter setelah dibersihkan: {len(cleaned)}")
    print("\n--- 500 karakter pertama ---")
    print(cleaned[:500])
    print("\n--- 500 karakter terakhir ---")
    print(cleaned[-500:])
